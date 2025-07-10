
import requests
import json
import time
from datetime import datetime

import os
from bs4 import BeautifulSoup
import re

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import requests
import io
from PIL import Image

# 获取当前脚本所在的目录路径
current_dir = os.path.dirname(os.path.abspath(__file__))
# 构造 params.json 的完整路径
json_file_path = os.path.join(current_dir, "params.json")

with open(json_file_path, "r", encoding="utf-8") as file:
    params_data = json.load(file)
    
    print(params_data['好房档案']['fakeid'])

def create_safe_dirname(project_name, publish_date):
    """创建安全的文件夹名称"""
    # 移除特殊字符
    project_name = re.sub(r'[\\/*?:"<>|]', "", project_name)
    publish_date = re.sub(r'[\\/*?:"<>|]', "", publish_date)
    # 合并为文件夹名
    dirname = f"{publish_date[:10]}_{project_name}"  # 只取日期部分
    return dirname[:180].replace('、','_')  # 限制长度防止路径过长

main_url = "https://mp.weixin.qq.com/cgi-bin/appmsgpublish"

headers = {
    "Authority": "mp.weixin.qq.com",
    "Accept": "*/*",
    "Accept-Encoding": "gzip, deflate, br, zstd",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "no-cache",
    "Pragma": "no-cache",
    "Priority": "u=1, i",
    "Referer": "https://mp.weixin.qq.com/cgi-bin/appmsg?t=media/appmsg_edit_v2&action=edit&isNew=1&type=77&createType=0&token=1825166898&lang=zh_CN&timestamp=1750328917785",
    "Sec-Ch-Ua": '"Google Chrome";v="137", "Chromium";v="137", "Not/A)Brand";v="24"',
    "Sec-Ch-Ua-Mobile": "?0",
    "Sec-Ch-Ua-Platform": '"Windows"',
    "Sec-Fetch-Dest": "empty",
    "Sec-Fetch-Mode": "cors",
    "Sec-Fetch-Site": "same-origin",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/137.0.0.0 Safari/537.36",
    "X-Requested-With": "XMLHttpRequest"
}

# cookies、token、fakeid，保存下来，这三者缺一不可

params = {
    "sub": "list",
    "search_field": "null",
    "begin": "0",
    "count": "5",
    "query": "",
    "fakeid": params_data['丁辰灵']['fakeid'],
    "fakeid": "MjM5Njg1OTQ3Mg==",
    "type": "101_1",
    "free_publish_type": "1",
    "sub_action": "list_ex",
    "fingerprint": "038dd85112e63c99d3900811a884a257",
    "token": params_data['token'],
    "lang": "zh_CN",
    "f": "json",
    "ajax": "1"
}

cookies=params_data['cookies']

response = requests.get(
    main_url,
    headers=headers,
    cookies=cookies,
    params=params
)

parsed_data = json.loads(response.text)

publish_page_str = parsed_data["publish_page"]
publish_page_data = json.loads(publish_page_str)
print(publish_page_data['total_count'])
masssend_count = publish_page_data['masssend_count']

for i in range(0, masssend_count, 5):
    print(i)
    print(f"正在处理第{i//5 + 1}页数据...")
    
    params['begin'] = i

    response = requests.get(
        main_url,
        headers=headers,
        cookies=cookies,
        params=params
    )

    parsed_data = json.loads(response.text)

    publish_page_str = parsed_data["publish_page"]
    publish_page_data = json.loads(publish_page_str)

    publish_list = publish_page_data['publish_list']
    for j in range(len(publish_list)):
        # 安全获取并解析 JSON
        publish_info = publish_list[j].get('publish_info', '')
        if not publish_info.strip():
            print(f"Skipping empty data at index {j}")
            continue

        try:
            single_data = json.loads(publish_info)
        except json.JSONDecodeError:
            print(f"Skipping invalid JSON at index {j}")
            continue

        create_time = single_data['appmsgex'][0]['create_time']
        dt = datetime.fromtimestamp(create_time)
        publish_date = dt.strftime("%Y年%m月%d日")
        publish_year = dt.year
        publish_month = dt.month
        publish_day = dt.day

        # 如果日期小于2025年，则结束爬虫
        if publish_year < 2025:
            break

        # print(single_data['sent_status']['total'])
        title = single_data['appmsgex'][0]['title']
        cover = single_data['appmsgex'][0]['cover']
        link = single_data['appmsgex'][0]['link']

        illegal_chars = [" ", "/", "\\", "|", '"', ':', '*', '?', '<', '>', '|', '\t']
        for char in illegal_chars:
            title = title.replace(char, "_")

        # 创建附图文件夹
        # base_output_dir = r'Y:\GOA-项目公示数据\公众号\建筑芝士'
        safe_dirname = create_safe_dirname(title, publish_date)
        # project_dir = os.path.join(base_output_dir, safe_dirname)
        # image_folder = os.path.join(project_dir, '附图')

        # text_path = os.path.join(project_dir, '内容.txt')

        # 获取网页内容
        url = link
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }

        response = requests.get(url, headers=headers)
        response.encoding = 'utf-8'
        soup = BeautifulSoup(response.text, 'html.parser')

        # 提取文字内容
        content_div = soup.find('div', id='img-content')
        if content_div:
            # 保存文字内容到word文件
            output_path = f'F:\\公众号\\丁辰灵\\{publish_date}_{title}.docx'
            output_path_img = f'F:\\公众号\\丁辰灵\\{publish_date}_{title}_img.docx'
                    
            if os.path.exists(output_path_img):
                continue  # 任一文件存在则跳过

            # if not os.path.exists(image_folder):
            #     os.makedirs(image_folder)

            # 创建Word文档
            doc = Document()
            doc_img = Document()
                
            # 查找指定的div或使用整个body
            content_div = soup.find('div', id='img-content')
            content = content_div if content_div else soup.body
            
            # 处理所有元素
            for element in content.find_all(recursive=True):
                # 处理文本
                if element.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    doc.add_paragraph(element.get_text())
                    doc_img.add_paragraph(element.get_text())
                
                # 处理图片
                elif element.name == 'img':
                    img_url = element.get('data-src') or element.get('src')
                    if img_url:
                        try:
                            response = requests.get(img_url, stream=True)
                            if response.status_code == 200:
                                image = Image.open(io.BytesIO(response.content))
                                image_path = 'temp_image.jpg'
                                image.save(image_path)
                                doc_img.add_picture(image_path, width=Inches(4))
                        except:
                            pass  # 如果图片无法加载，跳过
            
            # 保存Word文档
            doc.save(output_path)
            doc_img.save(output_path_img)
            print(f"成功保存Word文件到: {output_path}")
            print(f"成功保存Word文件到: {output_path_img}")
            
        else:
            print('未找到id为img-content的div')

        print('处理完成')
        time.sleep(3)







