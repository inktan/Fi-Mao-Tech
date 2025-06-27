
import requests
import json
import time
from datetime import datetime

import os
from bs4 import BeautifulSoup
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import re

def create_safe_dirname(project_name, publish_date):
    """创建安全的文件夹名称"""
    # 移除特殊字符
    project_name = re.sub(r'[\\/*?:"<>|]', "", project_name)
    publish_date = re.sub(r'[\\/*?:"<>|]', "", publish_date)
    # 合并为文件夹名
    dirname = f"{project_name}_{publish_date[:10]}"  # 只取日期部分
    return dirname[:180]  # 限制长度防止路径过长

main_url = "https://mp.weixin.qq.com/cgi-bin/appmsgpublish"

headers = {
    "Authority": "mp.weixin.qq.com",
    "Accept": "*/*",
    "Accept-Encoding": "gzip, deflate, br, zstd",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "no-cache",
    "Pragma": "no-cache",
    "Priority": "u=1, i",
    "Referer": "https://mp.weixin.qq.com/cgi-bin/appmsg?t=media/appmsg_edit_v2&action=edit&isNew=1&type=77&createType=0&token=1825166898&lang=zh_CN&timestamp=1750328917785",
    "Sec-Ch-Ua": '"Google Chrome";v="137", "Chromium";v="137", "Not/A)Brand";v="24"',
    "Sec-Ch-Ua-Mobile": "?0",
    "Sec-Ch-Ua-Platform": '"Windows"',
    "Sec-Fetch-Dest": "empty",
    "Sec-Fetch-Mode": "cors",
    "Sec-Fetch-Site": "same-origin",
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/137.0.0.0 Safari/537.36",
    "X-Requested-With": "XMLHttpRequest"
}

# cookies、token、fakeid，保存下来，这三者缺一不可

cookies = {
    "RK": "fCX0tqueWs",
    "ptcz": "7a473fcf51697bfc8b5f60a985ce8d0b00414a9e51fb375a399958d9fd27688a",
    "ua_id": "HMV7IE3np8vcDyIHAAAAAIMskFxnuPnsMuGxJzH7XPc=",
    "wxuin": "40479797108846",
    "mm_lang": "zh_CN",
    "pgv_pvid": "1704915140075719",
    "_t_qbtool_uid": "aaaaz2i556g4ttu1me7sikfqy0mp88cb",
    "_ga": "GA1.1.28274720.1743177134",
    "_ga_TPFW0KPXC1": "GS1.1.1743302505.2.0.1743302508.0.0.0",
    "fqm_pvqid": "adae4afe-e650-4e3c-9495-1592eb46efef",
    "b-user-id": "288c6bfa-76e7-63a1-0050-8b3cf7dba4f1",
    "rewardsn": "",
    "wxtokenkey": "777",
    "poc_sid": "HH7dU2ijlQ3OrFZNKco8U_N2roQ1dhWy6lPP2jPY",
    "_clck": "hwsk3d|1|fww|0",
    "uuid": "80372cee90a3620ec5d86e584e93f982",
    "rand_info": "CAESIEJabpaOLuLdqP1lKj6OCvh6MI8hOekRagQ5f5/j6cjI",
    "slave_bizuin": "3891696458",
    "data_bizuin": "3891696458",
    "bizuin": "3891696458",
    "data_ticket": "yKc64NjfsvUfqI5xB40bVWwrm8MII+tFPWwC9vDqtUjQi7GqUesf72T/RZcvHflF",
    "slave_sid": "blBnb1ZUZUtTeFFaeHdwb1JyOUhJamowS19sb3ZzbVlTSWhDc242N0NudDdJWnRnTVVBN1FFWFZ2bDF4bE0xaUtVQ1duN2hmbVZETjlvZE1oa2FobzFBeFBlTzBtaFZsOWk4bVA4TjJvTEV4OWFYTXJkYWVrekphbmQ5cnlFVmZPVzhhUVJCMHQyenVNTTFW",
    "slave_user": "gh_08417651a0c6",
    "xid": "fa399642748836c8c79799fdd487db89",
    "_clsk": "1h22fl8|1750328919743|3|1|mp.weixin.qq.com/weheat-agent/payload/record"
}

params = {
    "sub": "list",
    "search_field": "null",
    "begin": "0",
    "count": "5",
    "query": "",
    "fakeid": "MjM5NzAwMjk4Mw==",
    "type": "101_1",
    "free_publish_type": "1",
    "sub_action": "list_ex",
    "fingerprint": "038dd85112e63c99d3900811a884a257",
    "token": "1825166898",
    "lang": "zh_CN",
    "f": "json",
    "ajax": "1"
}

response = requests.get(
    main_url,
    headers=headers,
    cookies=cookies,
    params=params
)

parsed_data = json.loads(response.text)

publish_page_str = parsed_data["publish_page"]
publish_page_data = json.loads(publish_page_str)
print(publish_page_data['total_count'])
masssend_count = publish_page_data['masssend_count']

for i in range(0, masssend_count, 5):
    print(i)
    print(f"正在处理第{i//5 + 1}页数据...")
    
    params['begin'] = i

    response = requests.get(
        main_url,
        headers=headers,
        cookies=cookies,
        params=params
    )

    parsed_data = json.loads(response.text)

    publish_page_str = parsed_data["publish_page"]
    publish_page_data = json.loads(publish_page_str)

    for i in range(len(publish_page_data['publish_list'])):
        single_data = json.loads(publish_page_data['publish_list'][i]['publish_info'])
        sent_time = single_data['sent_info']['time']
        dt = datetime.fromtimestamp(sent_time)
        publish_date = dt.strftime("%Y年%m月%d日")
        year = dt.year

        # print(single_data['sent_status']['total'])
        title = single_data['appmsgex'][0]['title']
        cover = single_data['appmsgex'][0]['cover']
        link = single_data['appmsgex'][0]['link']

        # 创建附图文件夹
        base_output_dir = r'Y:\GOA-项目公示数据\公众号\搜建筑'
        safe_dirname = create_safe_dirname(title, publish_date)
        project_dir = os.path.join(base_output_dir, safe_dirname)
        image_folder = os.path.join(project_dir, '附图')
        if not os.path.exists(image_folder):
            os.makedirs(image_folder)

        text_path = os.path.join(project_dir, '内容.docx')

        # 获取网页内容
        url = link
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers)
        response.encoding = 'utf-8'
        soup = BeautifulSoup(response.text, 'html.parser')
        # 提取文字内容
        content_div = soup.find('div', id='img-content')
        if content_div:
            # 创建 Word 文档
            doc = Document()
            # 遍历所有子元素（段落、图片等）
            for element in content_div.children:
                # 处理文本段落
                if element.name == 'p' and element.text.strip():
                    doc.add_paragraph(element.text.strip())
                
                # 处理图片
                elif element.name == 'img':
                    img_url = element.get('data-src') or element.get('src')
                    if img_url and img_url.startswith('http'):
                        try:
                            img_data = requests.get(img_url).content
                            img_path = f'temp_img.{img_url.split(".")[-1].split("?")[0]}'
                            
                            # 保存临时图片
                            with open(img_path, 'wb') as f:
                                f.write(img_data)
                            
                            # 插入 Word
                            doc.add_picture(img_path, width=Inches(5))  # 控制图片宽度
                            os.remove(img_path)  # 删除临时图片
                        except Exception as e:
                            print(f"下载图片失败: {img_url}, 错误: {e}")
                
        # 处理标题（h1-h6）
        elif element.name and re.match(r'^h[1-6]$', element.name):
            doc.add_heading(element.text.strip(), level=int(element.name[1]))
    
        # 保存 Word 文档
        doc.save('微信公众号文章.docx')
        print("Word 文档已保存：微信公众号文章.docx")

        print('处理完成')
        time.sleep(2)







