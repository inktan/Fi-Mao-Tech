import json

# 1. 打开并读取 JSON 文件
with open(r'Y:\GOA-AIGC\98-goaTrainingData\ArchOctopus\github\Fi-Mao-Tech\web-crawler\公众号\params.json', 'r', encoding='utf-8') as file:
    data = json.load(file)  # 返回字典或列表（取决于 JSON 结构）

# 2. 使用数据（示例）
print(data)  # 打印整个 JSON 数据
print(data.keys())  # 访问具体字段（假设是字典）
print(data.values())  # 访问具体字段（假设是字典）

#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
import pdfkit

def fetch_page(url):
    """获取网页内容"""
    headers = {'User-Agent': 'Mozilla/5.0 (compatible; YourBot/0.1)'}
    try:
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        return BeautifulSoup(response.content, "html.parser")
    except requests.exceptions.RequestException as e:
        print(f"Error fetching {url}: {e}")
        return None

def parse_article(soup):
    """解析文章内容"""
    article_data = {
        'title': soup.find('h1').text.strip() if soup.find('h1') else "No Title",
        'content': "\n".join([p.text.strip() for p in soup.find_all('p')])
    }
    return article_data

def save_to_word(article_data, filename):
    """将文章数据保存为Word文档"""
    doc = Document()
    doc.add_heading(article_data['title'], level=1)
    doc.add_paragraph(article_data['content'])
    doc.save(filename)
    print(f"Article saved to {filename}")

def save_to_pdf(article_data, filename):
    """将文章数据保存为PDF文档"""
    html_content = f"<h1>{article_data['title']}</h1><p>{article_data['content']}</p>"
    pdfkit.from_string(html_content, filename)
    print(f"Article saved to {filename}")

def main():
    base_url = "https://mp.weixin.qq.com/s/nPq_zeP2ReAeef8Kzov8yA"  # 替换为目标网站的URL
    soup = fetch_page(base_url)
    if soup:
        article_data = parse_article(soup)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        word_filename = f'F:\\公众号\\丁辰灵\\article_{timestamp}.docx'
        pdf_filename = f'F:\\公众号\\丁辰灵\\article_{timestamp}.pdf'
        save_to_word(article_data, word_filename)
        save_to_pdf(article_data, pdf_filename)

if __name__ == "__main__":
    main()

